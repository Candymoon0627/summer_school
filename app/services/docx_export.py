from pathlib import Path
from typing import ClassVar
from uuid import uuid4

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.core.config import get_settings
from app.schemas.lesson import StructuredLesson, TrilingualText


class DocxExportService:
    default_font = "Tahoma"
    language_labels: ClassVar[dict[str, str]] = {
        "th": "Thai",
        "ms": "Local Malay",
        "en": "English",
    }

    def export_lesson(
        self,
        lesson: StructuredLesson,
        filename_prefix: str = "lesson",
        *,
        language: str = "th",
    ) -> Path:
        if language not in self.language_labels:
            raise ValueError(f"Unsupported lesson export language: {language}")
        settings = get_settings()
        export_dir = settings.local_storage_dir / "generated_docx"
        export_dir.mkdir(parents=True, exist_ok=True)

        path = export_dir / f"{filename_prefix}_{language}_{uuid4().hex[:8]}.docx"
        document = Document()
        self._configure_fonts(document)
        document.add_heading(self._text(lesson.title_trilingual, lesson.title, language), level=1)
        document.add_paragraph(self._text(lesson.summary_trilingual, lesson.summary, language))

        document.add_heading("Teaching Objectives", level=2)
        self._add_language_list(document, lesson.teaching_objectives_trilingual, lesson.teaching_objectives, language)

        document.add_heading("Materials", level=2)
        self._add_language_list(document, lesson.materials_trilingual, lesson.materials, language)

        document.add_heading("Lesson Flow", level=2)
        if lesson.lesson_flow_trilingual:
            for step in lesson.lesson_flow_trilingual:
                document.add_paragraph(
                    f"{self._text(step.phase, step.phase.en, language)} "
                    f"({step.minutes or '?'} min)",
                    style="List Bullet",
                )
                document.add_paragraph(self._text(step.teacher_action, "", language))
                if step.student_action:
                    document.add_paragraph(self._text(step.student_action, "", language))
        else:
            for step in lesson.lesson_flow:
                document.add_paragraph(
                    f"{step.phase} ({step.minutes or '?'} min): {step.teacher_action}",
                    style="List Bullet",
                )

        document.add_heading("Local Examples", level=2)
        self._add_language_list(document, lesson.local_examples_trilingual, lesson.local_examples, language)

        document.add_heading("Key Terms", level=2)
        for term in lesson.key_terms_trilingual:
            document.add_paragraph(
                f"Thai: {term.term_th} | Local Malay Helper: {term.helper_ms} | "
                f"English: {term.helper_en}"
            )

        document.add_heading("Practice Questions", level=2)
        self._add_language_list(
            document,
            lesson.practice_questions_trilingual,
            lesson.practice_questions,
            language,
            numbered=True,
        )

        document.add_heading("Board Plan", level=2)
        document.add_paragraph(self._text(lesson.board_plan_trilingual, lesson.board_plan, language))

        document.add_heading("Low-resource Alternative", level=2)
        document.add_paragraph(
            self._text(lesson.low_resource_plan_trilingual, lesson.low_resource_plan, language)
        )

        document.add_heading("Safety Notes", level=2)
        document.add_paragraph(
            self._text(lesson.safety_notes_trilingual, lesson.safety_notes or "", language)
        )

        document.add_heading("AI Notice", level=2)
        document.add_paragraph(
            "This lesson plan was AI-assisted and should be reviewed by the teacher "
            "before classroom use."
        )

        document.save(path)
        return path

    def _text(self, text: TrilingualText | None, fallback: str, language: str) -> str:
        if text is None:
            return fallback
        return getattr(text, language)

    def _add_language_list(
        self,
        document: Document,
        items: list[TrilingualText],
        fallback: list[str],
        language: str,
        *,
        numbered: bool = False,
    ) -> None:
        style = "List Number" if numbered else "List Bullet"
        if not items:
            for item in fallback:
                document.add_paragraph(item, style=style)
            return
        for item in items:
            document.add_paragraph(getattr(item, language), style=style)

    def _configure_fonts(self, document: Document) -> None:
        for style_name in ["Normal", "Heading 1", "Heading 2", "List Bullet", "List Number"]:
            if style_name not in document.styles:
                continue
            style = document.styles[style_name]
            font = style.font
            font.name = self.default_font
            if style_name == "Normal":
                font.size = Pt(11)
            self._set_rfonts(style.element.rPr)

    def _set_rfonts(self, rpr) -> None:
        if rpr is None:
            return
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr.get_or_add_rFonts()
        for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
            rfonts.set(qn(f"w:{key}"), self.default_font)
